"""Export tailored resumes from Markdown to a premium, ATS-safe DOCX and PDF.

Single column on purpose: two-column / sidebar layouts get scrambled by ATS
parsers (Workday, Greenhouse, Lever, Ashby). This renderer keeps the document
100% machine-parsable while looking modern — accent-colored name/title, section
rules, refined typography. DOCX is the safest submission format; the text-based
PDF is for human reviewers (recruiter, LinkedIn).
"""
from __future__ import annotations

import platform
import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from fpdf import FPDF
from fpdf.enums import XPos, YPos

# ---------- palette (matches the web app brand) ----------
INK = (31, 41, 55)       # gray-800 — body text / name
ACCENT = (79, 70, 229)   # indigo-600 — title, section headers, rules
MUTED = (107, 114, 128)  # gray-500 — contact, dates
RULE = (210, 213, 222)   # light divider
INK_HEX = "1F2937"
ACCENT_HEX = "4F46E5"
MUTED_HEX = "6B7280"

# ---------- cross-platform Unicode fonts (regular, bold, italic) ----------
_FONT_SETS = {
    "Darwin": [
        ("/System/Library/Fonts/Supplemental/Arial.ttf",
         "/System/Library/Fonts/Supplemental/Arial Bold.ttf",
         "/System/Library/Fonts/Supplemental/Arial Italic.ttf"),
        ("/Library/Fonts/Arial.ttf", "/Library/Fonts/Arial Bold.ttf", "/Library/Fonts/Arial Italic.ttf"),
    ],
    "Windows": [
        ("C:/Windows/Fonts/arial.ttf", "C:/Windows/Fonts/arialbd.ttf", "C:/Windows/Fonts/ariali.ttf"),
    ],
    "Linux": [
        ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
         "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
         "/usr/share/fonts/truetype/dejavu/DejaVuSans-Oblique.ttf"),
        ("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
         "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
         "/usr/share/fonts/truetype/liberation/LiberationSans-Italic.ttf"),
    ],
}


def _resolve_fonts() -> tuple[str, str, str] | None:
    """First existing (regular, bold, italic) triple for this OS; bold/italic
    fall back to regular when absent. None -> use fpdf core font (latin-1)."""
    for regular, bold, italic in _FONT_SETS.get(platform.system(), _FONT_SETS["Linux"]):
        if Path(regular).exists():
            reg = regular
            return reg, (bold if Path(bold).exists() else reg), (italic if Path(italic).exists() else reg)
    return None


BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
ITALIC_RE = re.compile(r"(?<!\*)\*([^*]+)\*(?!\*)")
CODE_RE = re.compile(r"`([^`]+)`")
NUMBERED_RE = re.compile(r"^\d+\.\s+")
ALL_BOLD_RE = re.compile(r"^\*\*(.+)\*\*$")


def _clean_inline(text: str) -> str:
    """Strip markdown the renderers don't handle (links -> 'text (url)', italics, code)."""
    text = LINK_RE.sub(lambda m: f"{m.group(1)} ({m.group(2)})", text)
    text = CODE_RE.sub(r"\1", text)
    text = ITALIC_RE.sub(r"\1", text)
    return text


# ---------- shared semantic parse ----------
# One markdown pass -> a list of (kind, text) blocks, so DOCX and PDF render the
# SAME structure. kinds: name, title, contact, hr, section, role, dates, bullet, para.
def _parse_blocks(md_text: str) -> list[tuple[str, str]]:
    blocks: list[tuple[str, str]] = []
    prev = None
    for raw in md_text.splitlines():
        line = raw.rstrip()
        s = line.strip()
        if not s:
            continue
        if s == "---":
            blocks.append(("hr", "")); prev = "hr"; continue
        if line.startswith("# "):
            blocks.append(("name", line[2:].strip())); prev = "name"; continue
        if line.startswith("## "):
            blocks.append(("section", line[3:].strip())); prev = "section"; continue
        if line.startswith("### "):
            blocks.append(("role", line[4:].strip())); prev = "role"; continue
        if line.startswith(("- ", "* ")):
            blocks.append(("bullet", line[2:].strip())); prev = "bullet"; continue
        if NUMBERED_RE.match(s):
            blocks.append(("bullet", NUMBERED_RE.sub("", s))); prev = "bullet"; continue
        m = ALL_BOLD_RE.match(s)
        if m:
            if prev == "name":
                blocks.append(("title", m.group(1).strip())); prev = "title"; continue
            if prev == "role":
                blocks.append(("dates", m.group(1).strip())); prev = "dates"; continue
            blocks.append(("para", s)); prev = "para"; continue
        # plain line: contact while still in the header block, else body paragraph
        if prev in ("name", "title", "contact"):
            blocks.append(("contact", s)); prev = "contact"; continue
        blocks.append(("para", s)); prev = "para"; continue
    return blocks


# ============================================================
#  DOCX
# ============================================================
def _set_border(paragraph, *, color: str, size: int) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    edge = OxmlElement("w:bottom")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(size))
    edge.set(qn("w:space"), "3")
    edge.set(qn("w:color"), color)
    pbdr.append(edge)
    pPr.append(pbdr)


def _runs(paragraph, text: str, *, bold=False, italic=False, size=None, color=None):
    """Add **bold**-aware runs to a paragraph with shared formatting."""
    text = _clean_inline(text)
    pieces, pos = [], 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            pieces.append((text[pos:m.start()], False))
        pieces.append((m.group(1), True))
        pos = m.end()
    if pos < len(text):
        pieces.append((text[pos:], False))
    if not pieces:
        pieces = [(text, False)]
    for chunk, is_bold in pieces:
        r = paragraph.add_run(chunk)
        r.bold = bold or is_bold
        r.italic = italic
        if size:
            r.font.size = Pt(size)
        if color:
            r.font.color.rgb = RGBColor(*color)


def markdown_to_docx(md_path: Path, out_path: Path) -> Path:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(*INK)
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(40)
        section.left_margin = section.right_margin = Pt(50)

    def para(space_before=0.0, space_after=2.0):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.06
        return p

    for kind, text in _parse_blocks(md_path.read_text(encoding="utf-8")):
        if kind == "name":
            p = para(0, 1)
            _runs(p, text, bold=True, size=21, color=INK)
        elif kind == "title":
            p = para(0, 3)
            _runs(p, text.upper(), bold=True, size=11.5, color=ACCENT)
        elif kind == "contact":
            p = para(0, 1)
            _runs(p, text, size=9, color=MUTED)
        elif kind == "hr":
            continue  # section rules carry the visual structure; skip raw dividers
        elif kind == "section":
            p = para(11, 4)
            _runs(p, text.upper(), bold=True, size=11.5, color=ACCENT)
            _set_border(p, color=ACCENT_HEX, size=6)
        elif kind == "role":
            p = para(6, 0)
            _runs(p, text, bold=True, size=10.5, color=INK)
        elif kind == "dates":
            p = para(0, 2)
            _runs(p, text, italic=True, size=9, color=MUTED)
        elif kind == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(1.5)
            p.paragraph_format.line_spacing = 1.06
            _runs(p, text, size=10.5, color=INK)
        else:  # para
            p = para(0, 3)
            _runs(p, text, size=10.5, color=INK)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


# ============================================================
#  PDF
# ============================================================
class _ResumePDF(FPDF):
    """Single-column, text-based PDF — premium look, ATS-safe.

    Uses multi_cell (not write) so wrapping breaks cleanly BETWEEN lines and
    page breaks never split a word/bullet mid-sentence; headers get orphan
    protection so a section/role title never sits alone at a page bottom.
    """

    def __init__(self):
        super().__init__(format="A4")
        self.set_margins(17, 13, 17)
        self.set_auto_page_break(auto=True, margin=14)
        fonts = _resolve_fonts()
        if fonts:
            reg, bold, ital = fonts
            self.add_font("Body", "", reg)
            self.add_font("Body", "B", bold)
            self.add_font("Body", "I", ital)
            self._fam = "Body"
        else:
            self._fam = "Helvetica"  # core latin-1 fallback when no system TTF exists
        self.add_page()

    def ensure(self, mm: float) -> None:
        """Orphan control: start a new page if less than `mm` remains."""
        if (self.h - self.b_margin) - self.get_y() < mm:
            self.add_page()

    def rule(self, color=RULE, width=0.4, gap=1.2) -> None:
        self.set_draw_color(*color)
        self.set_line_width(width)
        y = self.get_y() + gap
        self.line(self.l_margin, y, self.w - self.r_margin, y)
        self.set_y(y + gap)

    def block(self, text: str, size: float, *, style: str = "", color=INK,
              before: float = 0.0, after: float = 1.0, bullet: bool = False) -> None:
        """Render one markdown block via multi_cell (clean line-level wrapping)."""
        text = _clean_inline(text)
        cw = self.w - self.l_margin - self.r_margin
        lh = size * 0.52
        if before:
            self.ln(before)
        if bullet:
            self.set_x(self.l_margin)
            self.set_text_color(*ACCENT)
            self.set_font(self._fam, "B", size)
            self.cell(4.5, lh, "•")
            self.set_text_color(*color)
            self.set_font(self._fam, style, size)
            self.multi_cell(cw - 4.5, lh, text, markdown=True,
                            new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        else:
            self.set_text_color(*color)
            self.set_font(self._fam, style, size)
            self.multi_cell(cw, lh, text, markdown=True,
                            new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        if after:
            self.ln(after)


def markdown_to_pdf(md_path: Path, out_path: Path) -> Path:
    pdf = _ResumePDF()
    for kind, text in _parse_blocks(md_path.read_text(encoding="utf-8")):
        if kind == "name":
            pdf.block(text, 20, style="B", color=INK, after=0.6)
        elif kind == "title":
            pdf.block(text.upper(), 11, style="B", color=ACCENT, after=1.0)
        elif kind == "contact":
            if "ats match" in text.lower():
                pdf.block(text, 9.2, style="B", color=ACCENT, after=1.0)
            else:
                pdf.block(text, 8.6, color=MUTED, after=0.5)
        elif kind == "hr":
            pdf.rule(color=ACCENT, width=0.5)
        elif kind == "section":
            pdf.ensure(26)  # keep header + first line together
            pdf.block(text.upper(), 11, style="B", color=ACCENT, before=2.4, after=0.2)
            pdf.rule(color=ACCENT, width=0.4, gap=0.7)
            pdf.ln(0.8)
        elif kind == "role":
            pdf.ensure(22)
            pdf.block(text, 10.3, style="B", color=INK, before=1.6, after=0.2)
        elif kind == "dates":
            pdf.block(text, 8.8, style="I", color=MUTED, after=0.8)
        elif kind == "bullet":
            pdf.block(text, 9.4, color=INK, bullet=True, after=0.6)
        elif ALL_BOLD_RE.match(text.strip()):  # per-company project sub-header
            pdf.ensure(18)
            pdf.block(text, 9.6, color=INK, before=1.2, after=0.4)
        else:  # normal paragraph (e.g., summary)
            pdf.block(text, 9.5, color=INK, before=0.4, after=0.9)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    pdf.output(str(out_path))
    return out_path
