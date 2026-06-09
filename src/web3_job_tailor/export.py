"""Stage 8 — export ATS-safe outputs + lightweight QA.

Renders the tailored markdown to a single-column DOCX (legacy ATS parse DOCX
best) and writes the markdown alongside. QA warns on tables/images that scramble
ATS parsers. PDF export is optional (WeasyPrint) and left to the caller.
"""

from __future__ import annotations

import re
from pathlib import Path

from .models import JobRequirements, ResumeProfile, TailoredResume

_HEADINGS = {
    "en": ("Summary", "Skills", "Experience", "Education", "Certifications", "Languages"),
    "pt": ("Resumo", "Competências", "Experiência Profissional", "Formação Acadêmica",
           "Certificações", "Idiomas"),
}


def tailored_to_markdown(
    tailored: TailoredResume, profile: ResumeProfile, lang: str = "en"
) -> str:
    """Deterministic markdown view of a TailoredResume (DOCX source + QA input)."""
    h_summary, h_skills, h_exp, h_edu, h_cert, h_lang = _HEADINGS.get(lang, _HEADINGS["en"])
    c = profile.contact
    contact_bits = [b for b in (c.location, c.email, c.phone, c.linkedin, c.github) if b]

    lines: list[str] = [f"# {c.full_name}", ""]
    if tailored.headline:
        lines += [f"**{tailored.headline}**", ""]
    if contact_bits:
        lines += [" | ".join(contact_bits), ""]
    if tailored.summary:
        lines += [f"## {h_summary}", "", tailored.summary, ""]
    if tailored.skills:
        lines += [f"## {h_skills}", ""]
        lines += [f"- **{s.label}:** {s.details}" for s in tailored.skills]
        lines += [""]
    if tailored.experiences:
        lines += [f"## {h_exp}", ""]
        for exp in tailored.experiences:
            period = " - ".join(x for x in (exp.start, exp.end) if x)
            head = f"### {exp.title} | {exp.company}"
            lines += [head]
            meta = " | ".join(x for x in (period, exp.location) if x)
            if meta:
                lines += [meta]
            lines += [f"- {b.text}" for b in exp.bullets]
            lines += [""]
    if profile.education:
        lines += [f"## {h_edu}", ""]
        for edu in profile.education:
            year = f" ({edu.year})" if edu.year else ""
            lines += [f"- {edu.degree} — {edu.institution}{year}"]
        lines += [""]
    if profile.certifications:
        lines += [f"## {h_cert}", ""]
        lines += [f"- {cert}" for cert in profile.certifications]
        lines += [""]
    if profile.languages:
        lines += [f"## {h_lang}", ""]
        lines += [f"- {lang_}" for lang_ in profile.languages]
        lines += [""]
    return "\n".join(lines)


def ats_selfcheck_pdf(pdf_path: str | Path, must_contain: list[str]) -> list[str]:
    """Re-extract text from the generated PDF and return tokens NOT found.

    Phase 0 deliverable: do not trust the renderer's ATS claim — verify that the
    PDF's text layer actually carries the name/skills/companies we rendered.
    """
    import pymupdf  # bundled with pymupdf4llm

    doc = pymupdf.open(str(pdf_path))
    text = "".join(page.get_text() for page in doc).lower()
    return [t for t in must_contain if t and t.lower() not in text]

_BOLD = re.compile(r"\*\*(.+?)\*\*")
_ITALIC = re.compile(r"\*(.+?)\*")
_CODE = re.compile(r"`(.+?)`")
_LINK = re.compile(r"\[(.+?)\]\((.+?)\)")


def _strip_md(text: str) -> str:
    text = _LINK.sub(r"\1 (\2)", text)
    text = _BOLD.sub(r"\1", text)
    text = _ITALIC.sub(r"\1", text)
    text = _CODE.sub(r"\1", text)
    return text


def to_markdown_file(markdown: str, out_path: str | Path) -> Path:
    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(markdown, encoding="utf-8")
    return out


def to_docx(markdown: str, out_path: str | Path) -> Path:
    from docx import Document  # python-docx

    doc = Document()
    for raw in markdown.splitlines():
        line = raw.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(_strip_md(line[2:].strip()), level=0)
        elif line.startswith("## "):
            doc.add_heading(_strip_md(line[3:].strip()), level=1)
        elif line.startswith("### "):
            doc.add_heading(_strip_md(line[4:].strip()), level=2)
        elif line.startswith(("- ", "* ")):
            doc.add_paragraph(_strip_md(line[2:].strip()), style="List Bullet")
        else:
            doc.add_paragraph(_strip_md(line))
    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


def qa(markdown: str, requirements: JobRequirements) -> list[str]:
    """ATS hygiene checks on the generated markdown."""
    warnings: list[str] = []
    if re.search(r"\|.+\|", markdown):
        warnings.append("Possible table detected — ATS parsers may scramble it; use single-column text.")
    if "![" in markdown:
        warnings.append("Image detected — remove it; ATS cannot read images.")
    low = markdown.lower()
    missing = [s for s in requirements.hard_skills if s and s.lower() not in low]
    if missing:
        warnings.append(
            "Hard skills required by the job and absent from the output (real gaps, do not fabricate): "
            + ", ".join(missing)
        )
    return warnings
